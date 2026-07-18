from __future__ import annotations

import zipfile
from collections.abc import Iterable
from datetime import date, datetime
from importlib.metadata import PackageNotFoundError, version
from pathlib import Path
from typing import Any

from app.state.models import RawExtractedContent

"""本模块以只读方式解析 XLSX、DOCX 和文本型 PDF，并施加资源安全上限。"""


# 文档路由器明确允许解析的文件扩展名集合。
SUPPORTED_EXTENSIONS = frozenset({".xlsx", ".docx", ".pdf"})
# 默认允许读取的单个压缩文件或 PDF 的最大字节数。
DEFAULT_MAX_FILE_SIZE_BYTES = 100 * 1024 * 1024
# 默认允许 Office ZIP 容器声明的最大解压后总字节数。
DEFAULT_MAX_UNCOMPRESSED_SIZE_BYTES = 500 * 1024 * 1024
# 默认允许单个文档解析器返回的最大文本字符数。
DEFAULT_MAX_CHARACTERS = 5_000_000
# 默认允许 XLSX 解析器遍历的最大单元格数量。
DEFAULT_MAX_EXCEL_CELLS = 200_000
# 默认允许 PDF 解析器读取的最大页数。
DEFAULT_MAX_PDF_PAGES = 500


def _package_version(distribution_name: str) -> str:
    """返回解析依赖版本；依赖元数据不可用时返回 unknown。"""
    try:
        return version(distribution_name)
    except PackageNotFoundError:
        return "unknown"


def _validate_input_file(
    file_path: str | Path,
    expected_extensions: Iterable[str],
    max_file_size_bytes: int,
) -> Path:
    """校验解析输入是受支持、未超限且非符号链接的普通文件。"""
    if max_file_size_bytes <= 0:
        raise ValueError("max_file_size_bytes 必须大于零")

    original_path = Path(file_path).expanduser()
    if original_path.is_symlink():
        raise ValueError(f"为保证读取边界，拒绝解析符号链接：{original_path}")

    path = original_path.resolve(strict=True)
    if not path.is_file():
        raise ValueError(f"解析目标不是普通文件：{path}")

    normalized_extensions = {
        item.lower() if item.startswith(".") else f".{item.lower()}"
        for item in expected_extensions
    }
    if path.suffix.lower() not in normalized_extensions:
        raise ValueError(
            f"文件扩展名 {path.suffix.lower()} 不属于预期类型："
            f"{sorted(normalized_extensions)}"
        )
    if path.stat().st_size > max_file_size_bytes:
        raise ValueError(
            f"文件大小超过安全上限 {max_file_size_bytes} 字节：{path}"
        )
    return path


def _validate_zip_container(
    file_path: Path,
    max_uncompressed_size_bytes: int,
    max_members: int = 20_000,
) -> None:
    """检查 Office ZIP 容器的成员数量和声明解压大小，降低压缩炸弹风险。"""
    if max_uncompressed_size_bytes <= 0 or max_members <= 0:
        raise ValueError("ZIP 安全上限必须大于零")

    try:
        with zipfile.ZipFile(file_path) as archive:
            members = archive.infolist()
            if len(members) > max_members:
                raise ValueError(f"ZIP 容器成员数量超过安全上限 {max_members}")
            total_size = sum(member.file_size for member in members)
            if total_size > max_uncompressed_size_bytes:
                raise ValueError(
                    "ZIP 容器声明的解压后大小超过安全上限 "
                    f"{max_uncompressed_size_bytes} 字节"
                )
    except zipfile.BadZipFile as exc:
        raise ValueError(f"文件不是有效的 Office ZIP 容器：{file_path}") from exc


def _stringify_cell_value(value: Any) -> str:
    """把 Excel 单元格值转换为稳定、可比较的文本形式。"""
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    return str(value).strip()


def parse_xlsx_document(
    file_path: str | Path,
    *,
    max_file_size_bytes: int = DEFAULT_MAX_FILE_SIZE_BYTES,
    max_uncompressed_size_bytes: int = DEFAULT_MAX_UNCOMPRESSED_SIZE_BYTES,
    max_cells: int = DEFAULT_MAX_EXCEL_CELLS,
    max_characters: int = DEFAULT_MAX_CHARACTERS,
) -> RawExtractedContent:
    """只读提取 XLSX 的工作表、单元格值和公式文本。

    该工具不会执行公式、宏或外部链接，不会保存工作簿，也不会修改源文件。
    为控制恶意或异常文件风险，它会校验文件大小、ZIP 声明解压大小、遍历
    单元格数和输出字符数。``.xls`` 与启用宏的格式不在当前支持范围内。

    Args:
        file_path: 调用方授权读取的本地 ``.xlsx`` 文件。
        max_file_size_bytes: 压缩文件本身允许的最大字节数。
        max_uncompressed_size_bytes: ZIP 声明解压总大小上限。
        max_cells: 最多遍历的单元格数量。
        max_characters: 最多返回的文本字符数。

    Returns:
        包含连续文本、工作表结构、空关键字段和解析警告的记录。

    Raises:
        ImportError: 未安装 ``openpyxl`` 时抛出。
        ValueError: 文件类型错误、资源超限或容器无效时抛出。
        OSError: 文件无法读取时由操作系统抛出。
    """
    if max_cells <= 0 or max_characters <= 0:
        raise ValueError("max_cells 和 max_characters 必须大于零")

    path = _validate_input_file(file_path, {".xlsx"}, max_file_size_bytes)
    _validate_zip_container(path, max_uncompressed_size_bytes)

    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ImportError("解析 XLSX 需要安装 openpyxl") from exc

    workbook = load_workbook(
        filename=path,
        read_only=True,
        data_only=False,
        keep_links=False,
    )
    lines: list[str] = []
    sheets: list[dict[str, Any]] = []
    warnings: list[str] = []
    visited_cells = 0
    extracted_characters = 0
    truncated = False

    try:
        for worksheet in workbook.worksheets:
            sheet_info: dict[str, Any] = {
                "name": worksheet.title,
                "max_row": worksheet.max_row,
                "max_column": worksheet.max_column,
                "non_empty_cells": 0,
                "formula_cells": 0,
            }
            lines.append(f"[工作表: {worksheet.title}]")

            for row in worksheet.iter_rows():
                row_values: list[str] = []
                for cell in row:
                    visited_cells += 1
                    if visited_cells > max_cells:
                        truncated = True
                        break

                    value = _stringify_cell_value(cell.value)
                    row_values.append(value)
                    if value:
                        sheet_info["non_empty_cells"] += 1
                        if isinstance(cell.value, str) and cell.value.startswith("="):
                            sheet_info["formula_cells"] += 1

                if truncated:
                    break
                if any(row_values):
                    line = "\t".join(row_values).rstrip()
                    if extracted_characters + len(line) + 1 > max_characters:
                        truncated = True
                        break
                    lines.append(line)
                    extracted_characters += len(line) + 1

            sheets.append(sheet_info)
            if truncated:
                break
    finally:
        workbook.close()

    if truncated:
        warnings.append("内容达到安全提取上限，结果已截断")

    return RawExtractedContent(
        text="\n".join(lines),
        structure={
            "document_type": "xlsx",
            "parser": f"openpyxl/{_package_version('openpyxl')}",
            "sheet_count": len(workbook.sheetnames),
            "sheets": sheets,
            "visited_cells": min(visited_cells, max_cells),
            "truncated": truncated,
        },
        key_fields={},
        warnings=warnings,
    )


def parse_docx_document(
    file_path: str | Path,
    *,
    max_file_size_bytes: int = DEFAULT_MAX_FILE_SIZE_BYTES,
    max_uncompressed_size_bytes: int = DEFAULT_MAX_UNCOMPRESSED_SIZE_BYTES,
    max_characters: int = DEFAULT_MAX_CHARACTERS,
) -> RawExtractedContent:
    """只读提取 DOCX 的段落、标题和表格文本。

    该工具不会执行嵌入对象、宏或外部关系，也不会保存或修改源文档。
    它只处理调用方明确提供的 ``.docx``，并施加文件、ZIP 解压声明和输出
    字符数限制；旧版 ``.doc`` 不在当前支持范围内。

    Args:
        file_path: 调用方授权读取的本地 ``.docx`` 文件。
        max_file_size_bytes: 压缩文件本身允许的最大字节数。
        max_uncompressed_size_bytes: ZIP 声明解压总大小上限。
        max_characters: 最多返回的文本字符数。

    Returns:
        包含正文、标题/表格结构、空关键字段和解析警告的记录。

    Raises:
        ImportError: 未安装 ``python-docx`` 时抛出。
        ValueError: 文件类型错误、资源超限或容器无效时抛出。
        OSError: 文件无法读取时由操作系统抛出。
    """
    if max_characters <= 0:
        raise ValueError("max_characters 必须大于零")

    path = _validate_input_file(file_path, {".docx"}, max_file_size_bytes)
    _validate_zip_container(path, max_uncompressed_size_bytes)

    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("解析 DOCX 需要安装 python-docx") from exc

    document = Document(str(path))
    lines: list[str] = []
    heading_counts: dict[str, int] = {}
    extracted_characters = 0
    truncated = False

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if extracted_characters + len(text) + 1 > max_characters:
            truncated = True
            break
        lines.append(text)
        extracted_characters += len(text) + 1
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            heading_counts[style_name] = heading_counts.get(style_name, 0) + 1

    table_summaries: list[dict[str, int]] = []
    if not truncated:
        for table_index, table in enumerate(document.tables, start=1):
            lines.append(f"[表格: {table_index}]")
            row_count = 0
            max_columns = 0
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                line = "\t".join(values).rstrip()
                if extracted_characters + len(line) + 1 > max_characters:
                    truncated = True
                    break
                if line:
                    lines.append(line)
                    extracted_characters += len(line) + 1
                row_count += 1
                max_columns = max(max_columns, len(values))
            table_summaries.append({"rows": row_count, "columns": max_columns})
            if truncated:
                break

    warnings = ["内容达到安全提取上限，结果已截断"] if truncated else []
    return RawExtractedContent(
        text="\n".join(lines),
        structure={
            "document_type": "docx",
            "parser": f"python-docx/{_package_version('python-docx')}",
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "heading_counts": heading_counts,
            "tables": table_summaries,
            "truncated": truncated,
        },
        key_fields={},
        warnings=warnings,
    )


def parse_pdf_document(
    file_path: str | Path,
    *,
    max_file_size_bytes: int = DEFAULT_MAX_FILE_SIZE_BYTES,
    max_pages: int = DEFAULT_MAX_PDF_PAGES,
    max_characters: int = DEFAULT_MAX_CHARACTERS,
) -> RawExtractedContent:
    """只读提取文本型 PDF 的逐页文本和基础页面结构。

    该工具不执行 OCR、不渲染活动内容、不访问外部链接，也不修改 PDF。
    对扫描图片型 PDF，返回内容可能为空并附带警告；加密且无法用空密码解密的
    PDF 会明确失败，不会尝试猜测密码。

    Args:
        file_path: 调用方授权读取的本地 ``.pdf`` 文件。
        max_file_size_bytes: PDF 文件允许的最大字节数。
        max_pages: 最多读取的页数。
        max_characters: 最多返回的文本字符数。

    Returns:
        包含逐页文本、页数结构、空关键字段和解析警告的记录。

    Raises:
        ImportError: 未安装 ``pypdf`` 时抛出。
        ValueError: 文件类型错误、资源超限或加密文件无法解密时抛出。
        OSError: 文件无法读取时由操作系统抛出。
    """
    if max_pages <= 0 or max_characters <= 0:
        raise ValueError("max_pages 和 max_characters 必须大于零")

    path = _validate_input_file(file_path, {".pdf"}, max_file_size_bytes)
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError("解析 PDF 需要安装 pypdf") from exc

    reader = PdfReader(str(path), strict=False)
    if reader.is_encrypted and reader.decrypt("") == 0:
        raise ValueError("PDF 已加密，且无法使用空密码解密")
    if len(reader.pages) > max_pages:
        raise ValueError(f"PDF 页数超过安全上限 {max_pages}")

    lines: list[str] = []
    page_summaries: list[dict[str, Any]] = []
    warnings: list[str] = []
    extracted_characters = 0
    truncated = False

    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as exc:  # pypdf 对损坏页可能抛出多种解析异常。
            warnings.append(f"第 {page_number} 页文本提取失败：{exc}")
            text = ""

        page_summaries.append(
            {
                "page_number": page_number,
                "text_characters": len(text),
            }
        )
        page_block = f"[页面: {page_number}]\n{text}"
        if extracted_characters + len(page_block) + 1 > max_characters:
            truncated = True
            break
        lines.append(page_block)
        extracted_characters += len(page_block) + 1

    if truncated:
        warnings.append("内容达到安全提取上限，结果已截断")
    if not any(item["text_characters"] for item in page_summaries):
        warnings.append("PDF 未提取到文本，可能是扫描图片型文档；当前不执行 OCR")

    return RawExtractedContent(
        text="\n".join(lines),
        structure={
            "document_type": "pdf",
            "parser": f"pypdf/{_package_version('pypdf')}",
            "page_count": len(reader.pages),
            "pages": page_summaries,
            "truncated": truncated,
        },
        key_fields={},
        warnings=warnings,
    )


def parse_document(
    file_path: str | Path,
    **limits: int,
) -> RawExtractedContent:
    """按扩展名路由到支持的只读文档解析器。

    该工具只支持 ``.xlsx``、``.docx`` 和 ``.pdf``，不会根据内容猜测格式，
    不会执行宏、公式、OCR 或外部链接，也不会修改源文件。``limits`` 只应包含
    目标解析器公开的安全上限参数；未知参数会明确报错。

    Args:
        file_path: 调用方授权读取的本地文档路径。
        **limits: 传递给目标解析器的资源安全上限。

    Returns:
        对应解析器产生的统一原始内容记录。

    Raises:
        ValueError: 扩展名不受支持或安全上限不合法时抛出。
        TypeError: ``limits`` 包含目标解析器不接受的参数时抛出。
        ImportError: 对应解析依赖未安装时抛出。
        OSError: 文件无法读取时由操作系统抛出。
    """
    extension = Path(file_path).suffix.lower()
    if extension == ".xlsx":
        return parse_xlsx_document(file_path, **limits)
    if extension == ".docx":
        return parse_docx_document(file_path, **limits)
    if extension == ".pdf":
        return parse_pdf_document(file_path, **limits)
    raise ValueError(
        f"不支持的文件扩展名 {extension or '<empty>'}；"
        f"当前仅支持 {sorted(SUPPORTED_EXTENSIONS)}"
    )
