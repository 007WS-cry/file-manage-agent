from __future__ import annotations

import argparse
import hashlib
import json
from collections.abc import Sequence
from pathlib import Path
from typing import Any

from docx import Document

"""本脚本生成可重复的只读治理演示目录、请求信封、发送证据和 SHA-256 基线清单。"""


# 单次演示允许生成的最大业务文件数，避免误操作制造无界数据。
MAX_DEMO_FILES = 500

# 演示数据使用的固定逻辑时间，使重复生成的 JSON 元数据保持稳定。
DEMO_TIMESTAMP = "2026-07-27T00:00:00+08:00"


def build_argument_parser() -> argparse.ArgumentParser:
    """构建演示数据生成脚本的命令行参数解析器。

    Returns:
        包含输出目录和文件数量参数的解析器。
    """
    parser = argparse.ArgumentParser(
        description="生成 1.0.3 端到端演示所需的只读 DOCX 数据和请求清单。"
    )
    parser.add_argument(
        "--output-root",
        type=Path,
        default=Path(".artifacts/demo"),
        help="演示根目录；目录已存在且非空时拒绝覆盖。",
    )
    parser.add_argument(
        "--file-count",
        type=int,
        default=24,
        help=f"生成的 DOCX 数量，范围为 1 到 {MAX_DEMO_FILES}。",
    )
    return parser


def validate_empty_output_root(output_root: Path) -> Path:
    """校验输出根目录不会覆盖现有数据并返回规范化路径。

    Args:
        output_root: 用户显式指定的演示根目录。

    Returns:
        已解析的绝对目录路径。

    Raises:
        ValueError: 路径为符号链接、普通文件或非空目录时抛出。
    """
    candidate = output_root.expanduser()
    if candidate.is_symlink():
        raise ValueError("演示输出根目录不得是符号链接")
    resolved = candidate.resolve()
    if resolved.exists() and not resolved.is_dir():
        raise ValueError("演示输出根路径必须是目录")
    if resolved.exists() and any(resolved.iterdir()):
        raise ValueError(f"演示输出目录不是空目录，拒绝覆盖：{resolved}")
    return resolved


def validate_file_count(file_count: int) -> int:
    """校验演示文件数量处于受控范围。

    Args:
        file_count: 命令行传入的文件数量。

    Returns:
        已验证的文件数量。

    Raises:
        ValueError: 数量不是合法整数或超出安全范围时抛出。
    """
    if isinstance(file_count, bool) or not isinstance(file_count, int):
        raise ValueError("file_count 必须是整数")
    if file_count < 1 or file_count > MAX_DEMO_FILES:
        raise ValueError(f"file_count 必须位于 1 到 {MAX_DEMO_FILES} 之间")
    return file_count


def demo_group_label(index: int) -> str:
    """把序号编码为低相似度的稳定文档组标签。

    Args:
        index: 从零开始的文档组序号。

    Returns:
        由十六进制字符组成的固定长度组标签。
    """
    source = f"file-governance-demo-group:{index}".encode()
    return hashlib.sha256(source).hexdigest()[:12]


def create_demo_docx(path: Path, *, group_label: str, version: int) -> None:
    """创建一个不包含外部链接或宏的确定性 DOCX 演示文件。

    Args:
        path: 新 DOCX 文件的目标路径。
        group_label: 当前版本组的稳定标签。
        version: 当前文件的版本序号。
    """
    document = Document()
    document.core_properties.title = f"{group_label} 文件治理演示"
    document.add_heading(f"{group_label} 合同版本 {version}", level=1)
    document.add_paragraph(f"版本：{version}")
    document.add_paragraph(f"金额：CNY {1000 + version * 100}")
    document.add_paragraph(" ".join([group_label] * 30))
    document.save(path)


def sha256_file(path: Path) -> str:
    """流式计算普通文件的 SHA-256。

    Args:
        path: 等待计算摘要的普通文件。

    Returns:
        小写十六进制 SHA-256 字符串。
    """
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def build_input_manifest(input_root: Path) -> dict[str, Any]:
    """扫描输入目录并构造路径、大小和 SHA-256 基线。

    Args:
        input_root: 只读业务输入根目录。

    Returns:
        可写入 JSON 并用于运行前后不变性比较的清单。
    """
    files = []
    for path in sorted(item for item in input_root.rglob("*") if item.is_file()):
        files.append(
            {
                "relative_path": path.relative_to(input_root).as_posix(),
                "size_bytes": path.stat().st_size,
                "sha256": sha256_file(path),
            }
        )
    return {
        "schema_version": "1.0",
        "input_root": "input",
        "file_count": len(files),
        "files": files,
    }


def write_json(path: Path, payload: object) -> None:
    """以 UTF-8 和稳定缩进写入演示 JSON 文件。

    Args:
        path: JSON 输出路径。
        payload: 可以由标准库 JSON 编码的数据。
    """
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def build_request_envelope(output_root: Path, *, file_count: int) -> dict[str, Any]:
    """构造前台、后台和 Cron 可共用的治理请求信封。

    Args:
        output_root: 演示根目录的绝对路径。
        file_count: 允许扫描的输入文件上限。

    Returns:
        禁用外部模型且所有写路径均位于演示目录的请求信封。
    """
    return {
        "request": {
            "root_directory": str(output_root / "input"),
            "recursive": True,
            "allowed_extensions": [".docx"],
            "max_files": file_count,
            "grouping_similarity_threshold": 0.72,
            "auto_select_threshold": 0.82,
            "pdf_match_threshold": 0.82,
            "delivery_log_path": str(output_root / "delivery_log.json"),
            "use_llm_summary": False,
        },
        "workspace": {
            "input_root": str(output_root / "input"),
            "input_readonly": True,
            "artifact_root": str(output_root / "artifacts"),
            "report_root": str(output_root / "reports"),
        },
        "checkpoint": {
            "backend": "sqlite",
            "database_path": str(output_root / "checkpoints" / "foreground.sqlite3"),
        },
        "prompt": {"enabled": False},
        "hooks": {"enabled": False},
        "llm": {"enabled": False},
        "email_mcp": {"enabled": False},
    }


def generate_demo_data(output_root: Path, *, file_count: int) -> dict[str, Any]:
    """生成完整演示目录并返回可供终端展示的摘要。

    Args:
        output_root: 必须不存在或为空的演示根目录。
        file_count: 需要生成的 DOCX 文件数量。

    Returns:
        包含根目录、文件数量、请求和基线清单路径的摘要。
    """
    resolved_root = validate_empty_output_root(output_root)
    validated_count = validate_file_count(file_count)
    input_root = resolved_root / "input"
    input_root.mkdir(parents=True, exist_ok=True)
    for directory_name in ("artifacts", "reports", "database", "checkpoints"):
        (resolved_root / directory_name).mkdir()

    for index in range(validated_count):
        group_index = index // 2
        version = index % 2 + 1
        label = demo_group_label(group_index)
        file_name = f"{label}_v{version}.docx"
        create_demo_docx(
            input_root / file_name,
            group_label=label,
            version=version,
        )

    manifest = build_input_manifest(input_root)
    request_envelope = build_request_envelope(
        resolved_root,
        file_count=validated_count,
    )
    first_file = manifest["files"][0]
    delivery_log = {
        "schema_version": "1.0",
        "deliveries": [
            {
                "id": "demo-delivery-001",
                "attachment_name": Path(first_file["relative_path"]).name,
                "attachment_sha256": first_file["sha256"],
                "normalized_digest": None,
                "sent_at": DEMO_TIMESTAMP,
                "recipient_label": "demo-customer",
                "customer_confirmed": True,
                "evidence_ref": "local-log://demo/delivery-001",
            }
        ],
    }
    write_json(resolved_root / "manifest.before.json", manifest)
    write_json(resolved_root / "request.json", request_envelope)
    write_json(
        resolved_root / "background_submission.json",
        {
            "execution_mode": "background",
            "max_attempts": 3,
            "payload": request_envelope,
        },
    )
    write_json(resolved_root / "delivery_log.json", delivery_log)
    return {
        "status": "generated",
        "output_root": str(resolved_root),
        "file_count": validated_count,
        "request_path": str(resolved_root / "request.json"),
        "manifest_path": str(resolved_root / "manifest.before.json"),
    }


def main(argv: Sequence[str] | None = None) -> int:
    """解析参数、生成演示数据并向标准输出写入 JSON 摘要。

    Args:
        argv: 可选命令行参数序列；省略时读取当前进程参数。

    Returns:
        成功时返回零；参数或安全校验失败时由 argparse 返回非零。
    """
    arguments = build_argument_parser().parse_args(argv)
    try:
        summary = generate_demo_data(
            arguments.output_root,
            file_count=arguments.file_count,
        )
    except (OSError, ValueError) as error:
        build_argument_parser().error(str(error))
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
